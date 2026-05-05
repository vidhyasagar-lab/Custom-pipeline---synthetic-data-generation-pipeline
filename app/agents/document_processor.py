"""Agent 1: Semantic Chunker (LangGraph Node)
Loads multiple documents (.docx and .pdf) and applies semantic chunking —
splits by meaning boundaries using embedding similarity between consecutive
sentences. Produces chunks with rich metadata and stores chunk embeddings
in state for reuse by downstream agents.

Supports document-level caching: if a document has been chunked before
(same content hash), reuses cached chunks + embeddings.
"""

import asyncio
import hashlib
import json
import time
from pathlib import Path

import numpy as np
from docx import Document as DocxDocument

from app.core.config import get_settings
from app.core.llm import get_azure_embeddings
from app.core.logging_config import get_agent_logger
from app.core.graph_state import GraphState

logger = get_agent_logger("Agent1:SemanticChunker")


def _compute_file_hash(file_path: Path) -> str:
    """Compute SHA-256 hash of file content for caching."""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            h.update(block)
    return h.hexdigest()


def _get_cache_path(file_hash: str, cache_dir: Path) -> Path:
    """Get the cache file path for a given content hash."""
    return cache_dir / f"chunks_{file_hash}.json"


def _load_cached_chunks(file_hash: str, cache_dir: Path) -> tuple[list[dict], list[list[float]]] | None:
    """Load cached chunks + embeddings if available."""
    cache_path = _get_cache_path(file_hash, cache_dir)
    if not cache_path.exists():
        return None
    try:
        with open(cache_path, "r", encoding="utf-8") as f:
            cached = json.load(f)
        chunks = cached.get("chunks", [])
        embeddings = cached.get("embeddings", [])
        if chunks and embeddings and len(chunks) == len(embeddings):
            return chunks, embeddings
        return None
    except Exception:
        return None


def _save_chunks_to_cache(
    file_hash: str, cache_dir: Path,
    chunks: list[dict], embeddings: list[list[float]],
):
    """Save chunks + embeddings to cache."""
    try:
        cache_dir.mkdir(parents=True, exist_ok=True)
        cache_path = _get_cache_path(file_hash, cache_dir)
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump({"chunks": chunks, "embeddings": embeddings}, f)
    except Exception as e:
        logger.warning(f"Failed to save cache: {e}")


def load_docx(file_path: Path) -> tuple[str, list[dict]]:
    """Load a .docx file and extract text content with heading structure.
    
    Returns (full_text, heading_map) where heading_map is a list of
    {char_offset, heading_text, heading_level} for hierarchical chunking.
    """
    try:
        logger.set_step("load_docx")
        logger.info(f"Opening file: {file_path.name}")

        if not file_path.exists():
            raise FileNotFoundError(f"File does not exist: {file_path}")

        doc = DocxDocument(str(file_path))
        full_text = []
        heading_map = []  # P3.11: Track heading positions

        para_count = 0
        current_offset = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect headings by style name
                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.replace("Heading", "").strip())
                    except ValueError:
                        level = 1
                    heading_map.append({
                        "char_offset": current_offset,
                        "heading_text": text,
                        "heading_level": level,
                    })

                full_text.append(text)
                current_offset += len(text) + 2  # +2 for \n\n join
                para_count += 1

        logger.info(f"Extracted {para_count} paragraphs, {len(heading_map)} headings from document body")

        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    joined = " | ".join(row_text)
                    full_text.append(joined)
                    current_offset += len(joined) + 2
                    table_count += 1

        logger.info(f"Extracted {table_count} rows from tables")

        combined = "\n\n".join(full_text)
        if not combined.strip():
            raise ValueError(f"Document is empty or contains no extractable text: {file_path.name}")

        logger.info(f"Total document size: {len(combined):,} characters")
        return combined, heading_map

    except (FileNotFoundError, ValueError):
        raise
    except Exception as e:
        logger.error(f"Failed to load document '{file_path.name}': {e}")
        raise RuntimeError(f"Error loading document '{file_path.name}': {e}") from e


def load_pdf(file_path: Path) -> str:
    """Load a .pdf file and extract text content using pymupdf."""
    try:
        logger.set_step("load_pdf")
        logger.info(f"Opening PDF: {file_path.name}")

        if not file_path.exists():
            raise FileNotFoundError(f"File does not exist: {file_path}")

        import pymupdf

        doc = pymupdf.open(str(file_path))
        full_text = []

        for page_num, page in enumerate(doc):
            text = page.get_text().strip()
            if text:
                full_text.append(text)

        doc.close()

        logger.info(f"Extracted text from {len(full_text)} pages")

        combined = "\n\n".join(full_text)
        if not combined.strip():
            raise ValueError(f"PDF is empty or contains no extractable text: {file_path.name}")

        logger.info(f"Total document size: {len(combined):,} characters")
        return combined

    except (FileNotFoundError, ValueError):
        raise
    except ImportError:
        logger.error("pymupdf not installed. Install with: pip install pymupdf")
        raise RuntimeError("pymupdf is required for PDF support. Install with: pip install pymupdf")
    except Exception as e:
        logger.error(f"Failed to load PDF '{file_path.name}': {e}")
        raise RuntimeError(f"Error loading PDF '{file_path.name}': {e}") from e


def load_document(file_path: Path) -> tuple[str, list[dict]]:
    """Load a document based on its file extension (.docx or .pdf).
    
    Returns (text, heading_map). PDF returns empty heading_map.
    """
    ext = file_path.suffix.lower()
    if ext == ".docx":
        return load_docx(file_path)
    elif ext == ".pdf":
        return load_pdf(file_path), []
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .docx, .pdf")


def _assign_heading_paths(chunks: list[dict], heading_map: list[dict], raw_text: str):
    """Assign heading_path and parent_section to chunks based on document heading structure.
    
    Each chunk gets a heading_path like ["Chapter 1", "Section 1.2", "Subsection"] 
    based on which headings precede its content in the document.
    """
    if not heading_map:
        return

    for chunk in chunks:
        chunk_text = chunk.get("page_content", "")
        # Find approximate position of chunk in raw text
        chunk_pos = raw_text.find(chunk_text[:100]) if chunk_text else -1
        if chunk_pos < 0:
            continue

        # Build heading path: most recent heading at each level up to this position
        heading_stack: dict[int, str] = {}
        for h in heading_map:
            if h["char_offset"] > chunk_pos:
                break
            level = h["heading_level"]
            heading_stack[level] = h["heading_text"]
            # Clear deeper levels when a higher-level heading appears
            for deeper in list(heading_stack.keys()):
                if deeper > level:
                    del heading_stack[deeper]

        if heading_stack:
            path = [heading_stack[k] for k in sorted(heading_stack.keys())]
            chunk["metadata"]["heading_path"] = path
            chunk["metadata"]["parent_section"] = path[-1] if path else ""


def resolve_document_paths(state: dict) -> list[Path]:
    """Resolve one or more document paths from state or data directory."""
    try:
        settings = get_settings()

        doc_paths = state.get("document_paths", [])
        if doc_paths:
            paths = [Path(p) for p in doc_paths]
        else:
            doc_path = state.get("document_path", "")
            if doc_path:
                paths = [Path(doc_path)]
            else:
                paths = sorted(
                    list(settings.data_dir.glob("*.docx")) +
                    list(settings.data_dir.glob("*.pdf"))
                )

        if not paths:
            raise ValueError(f"No supported documents found in {settings.data_dir}")

        for p in paths:
            if not p.exists():
                raise FileNotFoundError(f"Document not found: {p}")

        return paths

    except (ValueError, FileNotFoundError):
        raise
    except Exception as e:
        logger.error(f"Failed to resolve document paths: {e}")
        raise RuntimeError(f"Error resolving document paths: {e}") from e


def split_into_sentences(text: str) -> list[str]:
    """Split text into sentences using NLTK for accurate splitting."""
    try:
        import nltk
        try:
            from nltk.tokenize import sent_tokenize
            sentences = sent_tokenize(text)
        except LookupError:
            nltk.download("punkt_tab", quiet=True)
            from nltk.tokenize import sent_tokenize
            sentences = sent_tokenize(text)

        # Filter out very short fragments
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 20]

        if not sentences:
            raise ValueError("No valid sentences found after splitting.")
        return sentences

    except ValueError:
        raise
    except Exception as e:
        logger.warning(f"NLTK sentence splitting failed ({e}), falling back to regex")
        import re
        raw = re.split(r'(?<=[.!?])\s+|\n\n+', text)
        sentences = [s.strip() for s in raw if s.strip() and len(s.strip()) > 20]
        if not sentences:
            raise ValueError("No valid sentences found after splitting.")
        return sentences


def compute_cosine_similarity(a: list[float], b: list[float]) -> float:
    """Compute cosine similarity between two vectors."""
    try:
        a_arr = np.array(a)
        b_arr = np.array(b)
        dot = np.dot(a_arr, b_arr)
        norm = np.linalg.norm(a_arr) * np.linalg.norm(b_arr)
        if norm == 0:
            return 0.0
        return float(dot / norm)
    except Exception as e:
        logger.error(f"Cosine similarity computation failed: {e}")
        return 0.0


async def semantic_chunk(
    sentences: list[str],
    source_file: str = "",
    doc_index: int = 0,
    global_chunk_offset: int = 0,
    similarity_threshold: float = 0.75,
    min_chunk_size: int = 200,
    max_chunk_size: int = 3000,
) -> tuple[list[dict], list[list[float]]]:
    """Perform semantic chunking. Returns (chunks, chunk_embeddings)."""
    try:
        logger.set_step("semantic_chunk")

        if not sentences:
            return [], []

        embeddings_model = get_azure_embeddings()

        logger.info(f"Embedding {len(sentences)} sentences from '{source_file}'...")
        batch_size = 50
        all_embeddings = []
        for i in range(0, len(sentences), batch_size):
            batch = sentences[i:i + batch_size]
            try:
                batch_embs = await embeddings_model.aembed_documents(batch)
                all_embeddings.extend(batch_embs)
                logger.info(f"  Embedded batch {i // batch_size + 1}/{(len(sentences) - 1) // batch_size + 1}")
            except Exception as e:
                logger.error(f"  Embedding batch {i // batch_size + 1} failed: {e}")
                raise RuntimeError(f"Embedding failed at batch {i // batch_size + 1}: {e}") from e

        # Compute similarities between consecutive sentences
        similarities = []
        for i in range(len(all_embeddings) - 1):
            sim = compute_cosine_similarity(all_embeddings[i], all_embeddings[i + 1])
            similarities.append(sim)

        avg_sim = np.mean(similarities) if similarities else 0
        logger.info(f"Average consecutive similarity: {avg_sim:.3f}")
        logger.info(f"Similarity threshold: {similarity_threshold}")

        # Group sentences into chunks by similarity breaks
        chunks = []
        chunk_sentence_embeddings = []
        current_sentences = [sentences[0]]
        current_emb_indices = [0]
        current_start = 0

        def _make_chunk(sents, start_idx, end_idx, emb_indices):
            text = " ".join(sents)
            local_id = len(chunks)
            # Average the sentence embeddings for this chunk
            chunk_emb = np.mean([all_embeddings[ei] for ei in emb_indices], axis=0).tolist()
            chunk_sentence_embeddings.append(chunk_emb)
            return {
                "page_content": text,
                "metadata": {
                    "chunk_id": global_chunk_offset + local_id,
                    "doc_chunk_id": local_id,
                    "source_file": source_file,
                    "doc_index": doc_index,
                    "sentence_start": start_idx,
                    "sentence_end": end_idx,
                    "num_sentences": len(sents),
                    "char_length": len(text),
                },
            }

        for i, sim in enumerate(similarities):
            current_text = " ".join(current_sentences)

            if sim < similarity_threshold and len(current_text) >= min_chunk_size:
                chunks.append(_make_chunk(current_sentences, current_start, i, current_emb_indices))
                current_sentences = [sentences[i + 1]]
                current_emb_indices = [i + 1]
                current_start = i + 1
            else:
                current_sentences.append(sentences[i + 1])
                current_emb_indices.append(i + 1)

                current_text = " ".join(current_sentences)
                if len(current_text) > max_chunk_size:
                    chunks.append(_make_chunk(current_sentences, current_start, i + 1, current_emb_indices))
                    current_sentences = []
                    current_emb_indices = []
                    current_start = i + 2

        if current_sentences:
            chunks.append(_make_chunk(current_sentences, current_start, len(sentences) - 1, current_emb_indices))

        if not chunks:
            raise ValueError(f"Semantic chunking produced 0 chunks for '{source_file}'.")

        return chunks, chunk_sentence_embeddings

    except (ValueError, RuntimeError):
        raise
    except Exception as e:
        logger.error(f"Semantic chunking failed for '{source_file}': {e}")
        raise RuntimeError(f"Semantic chunking error for '{source_file}': {e}") from e


async def semantic_chunker_node(state: GraphState) -> dict:
    """LangGraph node: Load one or more documents and perform semantic chunking."""
    try:
        with logger.phase("Semantic Chunking"):
            settings = get_settings()

            # --- Step 1: Resolve document paths ---
            with logger.step("Resolve documents"):
                doc_paths = resolve_document_paths(state)
                logger.info(f"Documents to process: {len(doc_paths)}")
                for p in doc_paths:
                    logger.info(f"  - {p.name} ({p.suffix})")

            # --- Step 2: Load & chunk each document (parallel) ---
            all_chunks = []
            all_chunk_embeddings = []
            all_raw_texts = {}
            total_sentences = 0
            failed_docs = []

            async def _process_single_doc(doc_idx: int, doc_path: Path, semaphore: asyncio.Semaphore) -> dict:
                """Process a single document with semaphore for API rate limiting.
                Uses content-hash caching to skip re-processing identical documents.
                """
                async with semaphore:
                    # Check cache first
                    file_hash = _compute_file_hash(doc_path)
                    cached = _load_cached_chunks(file_hash, settings.cache_dir)
                    if cached is not None:
                        doc_chunks, doc_embs = cached
                        logger.info(f"  {doc_path.name}: CACHE HIT ({len(doc_chunks)} chunks)")
                        raw_text, _ = load_document(doc_path)
                        return {
                            "doc_path": doc_path,
                            "raw_text": raw_text,
                            "sentences": sum(c["metadata"].get("num_sentences", 0) for c in doc_chunks),
                            "chunks": doc_chunks,
                            "embeddings": doc_embs,
                            "cached": True,
                        }

                    raw_text, heading_map = load_document(doc_path)
                    sentences = split_into_sentences(raw_text)
                    logger.info(f"  {doc_path.name}: {len(sentences)} sentences")

                    start = time.time()
                    doc_chunks, doc_embs = await semantic_chunk(
                        sentences,
                        source_file=doc_path.name,
                        doc_index=doc_idx,
                        global_chunk_offset=0,  # Will reindex after merge
                        similarity_threshold=settings.similarity_threshold,
                        min_chunk_size=200,
                        max_chunk_size=3000,
                    )
                    elapsed = time.time() - start

                    # Save to cache
                    _save_chunks_to_cache(file_hash, settings.cache_dir, doc_chunks, doc_embs)
                    logger.info(f"  {doc_path.name}: {len(sentences)} sentences -> {len(doc_chunks)} chunks in {elapsed:.1f}s (cached)")

                    # P3.11: Enrich chunks with heading hierarchy
                    if heading_map:
                        _assign_heading_paths(doc_chunks, heading_map, raw_text)

                    return {
                        "doc_path": doc_path,
                        "raw_text": raw_text,
                        "sentences": len(sentences),
                        "chunks": doc_chunks,
                        "embeddings": doc_embs,
                        "cached": False,
                    }

            if len(doc_paths) == 1:
                # Single document — no parallelism overhead
                try:
                    with logger.step(f"Process doc 1/1: {doc_paths[0].name}"):
                        sem = asyncio.Semaphore(1)
                        result = await _process_single_doc(0, doc_paths[0], sem)
                        all_raw_texts[result["doc_path"].name] = result["raw_text"]
                        total_sentences += result["sentences"]
                        all_chunks.extend(result["chunks"])
                        all_chunk_embeddings.extend(result["embeddings"])
                except Exception as e:
                    logger.error(f"Failed to process document '{doc_paths[0].name}': {e}")
                    failed_docs.append({"document": doc_paths[0].name, "error": str(e)})
            else:
                # Multiple documents — process in parallel with semaphore
                with logger.step(f"Process {len(doc_paths)} documents in parallel"):
                    # Limit concurrent embedding API calls (2 docs at a time)
                    semaphore = asyncio.Semaphore(min(len(doc_paths), 2))
                    tasks = [
                        _process_single_doc(idx, path, semaphore)
                        for idx, path in enumerate(doc_paths)
                    ]
                    results = await asyncio.gather(*tasks, return_exceptions=True)

                    for idx, result in enumerate(results):
                        if isinstance(result, Exception):
                            logger.error(f"Failed to process document '{doc_paths[idx].name}': {result}")
                            failed_docs.append({"document": doc_paths[idx].name, "error": str(result)})
                        else:
                            all_raw_texts[result["doc_path"].name] = result["raw_text"]
                            total_sentences += result["sentences"]
                            all_chunks.extend(result["chunks"])
                            all_chunk_embeddings.extend(result["embeddings"])

                    # Reindex chunk_ids globally after parallel merge
                    for global_idx, chunk in enumerate(all_chunks):
                        chunk["metadata"]["chunk_id"] = global_idx

            if not all_chunks:
                raise RuntimeError(
                    f"All documents failed to produce chunks. "
                    f"Failed: {[d['document'] for d in failed_docs]}"
                )

            # Validate chunk-embedding alignment
            if len(all_chunks) != len(all_chunk_embeddings):
                logger.error(
                    f"Chunk-embedding mismatch: {len(all_chunks)} chunks vs "
                    f"{len(all_chunk_embeddings)} embeddings. Truncating to shorter."
                )
                min_len = min(len(all_chunks), len(all_chunk_embeddings))
                all_chunks = all_chunks[:min_len]
                all_chunk_embeddings = all_chunk_embeddings[:min_len]

            # --- Step 3: Summary ---
            with logger.step("Summary"):
                chunk_sizes = [c["metadata"]["char_length"] for c in all_chunks]
                avg_size = sum(chunk_sizes) / len(chunk_sizes) if chunk_sizes else 0

                logger.info(f"Semantic chunking complete for {len(doc_paths)} document(s)")
                logger.info(f"  Total sentences: {total_sentences}")
                logger.info(f"  Total chunks: {len(all_chunks)}")
                logger.info(f"  Chunk embeddings stored for reuse: {len(all_chunk_embeddings)}")
                logger.info(f"  Avg chunk size: {avg_size:.0f} chars")
                if chunk_sizes:
                    logger.info(f"  Min: {min(chunk_sizes)}, Max: {max(chunk_sizes)}")
                if failed_docs:
                    logger.warning(f"  Failed documents: {len(failed_docs)}")

                doc_counts = {}
                for c in all_chunks:
                    src = c["metadata"]["source_file"]
                    doc_counts[src] = doc_counts.get(src, 0) + 1
                for doc_name, cnt in doc_counts.items():
                    logger.info(f"  {doc_name}: {cnt} chunks")

            errors = [f"Doc '{d['document']}': {d['error']}" for d in failed_docs]

            return {
                "raw_text": "\n\n".join(all_raw_texts.values()),
                "document_paths": [str(p) for p in doc_paths],
                "chunks": all_chunks,
                "chunk_embeddings": all_chunk_embeddings,
                "errors": errors,
                "phase_timings": [{
                    "phase": "semantic_chunking",
                    "documents_processed": len(doc_paths) - len(failed_docs),
                    "documents_failed": len(failed_docs),
                    "total_sentences": total_sentences,
                    "total_chunks": len(all_chunks),
                    "avg_chunk_size": round(avg_size),
                    "chunks_per_document": doc_counts,
                }],
            }

    except Exception as e:
        logger.error(f"Semantic chunker node failed: {e}")
        raise
